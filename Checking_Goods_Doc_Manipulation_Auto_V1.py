from docx import Document



import datetime



from selenium import webdriver



from selenium.webdriver.support.ui import WebDriverWait



from selenium.webdriver.support import expected_conditions



import time



from selenium.webdriver.support.select import Select



import os



import pyautogui



import pyperclip



class goods_checking_document_generator():

    def application_with_slash(self):

        return "{}/{}/{}/{}/{}".format(self.application_number[0:2], self.application_number[2:4], self.application_number[4:6], self.application_number[6:9], self.application_number[9:14])



    def application_number_check(self):

        if len(self.application_number) > 14 or len(self.application_number) < 12:

            return False

        else:

            return True





    def application_number_convert(self):

        if len(self.application_number) == 14:

            return self.application_number

        elif len(self.application_number) == 12:

            return self.application_number[0:2] + "  " + self.application_number[2:14]

        else:

            return ''





    def application_number_transformation(self):

        if self.application_number_check() == True:

            return self.application_number_convert()

        else:

            return ''





    def filling_application_number(self):

        application_number_transformation_finished = self.application_number_transformation() 

        self.driver.execute_script("sf1Clear();")

        search_input = self.driver.find_element_by_name("declNo1")

        search_input.send_keys(application_number_transformation_finished)





    def filling_item_number(self, number_input):

        search_input = self.driver.find_element_by_name("itemNo")

        search_input.clear()

        search_input.send_keys(number_input)



    def pressing_searching_button_F6(self):



        commit = ''

        commit = self.driver.find_element_by_name("btn_F6")

        while commit == '':

            time.sleep(0.5)

            commit = self.driver.find_element_by_name("btn_F6")



        commit.click()





    def pressing_next_item_buttom(self):

        click_lock = True

        while click_lock:

            try:

                searching_page_down_button = self.driver.find_element_by_id("rcfw")

                searching_page_down_button.click()

            except Exception as error:

                continue

            click_lock = False





    def collecting_element_by_javascript_commend(self, id_input):

        script_commend = "return $('#{}').val();".format(id_input)

        

        check_element_empty = ''

        check_element_empty = str(self.driver.execute_script(script_commend))

        while check_element_empty == '':

             time.sleep(1)

             check_element_empty = str(self.driver.execute_script(script_commend))

        return check_element_empty





    def collecting_element_by_finding_id(self, id_input):

        check_element_empty = ''

        select = Select(self.driver.find_element_by_id(id_input))

        check_element_empty = str(select.first_selected_option.text)

        while check_element_empty == '':

            time.sleep(1)

            select = Select(self.driver.find_element_by_id(id_input))

            check_element_empty = str(select.first_selected_option.text)



        return check_element_empty





    def checking_status_Message(self):

        click_lock = True

        while click_lock:

            try:

                searching_page_down_button = self.driver.find_element_by_id("statusMsg")

                searching_page_down_button.click()

            except Exception as error:

                continue

            click_lock = False



    def finding_total_page_number(self):



        total_page_number_xpath = "/html/body/div[2]/div[3]/form/div[1]/div[2]/table/tbody/tr/td/span[1]/span"

        total_page_number_element = str(self.driver.find_element_by_xpath(total_page_number_xpath).text)



        while total_page_number_element == '':

            time.sleep(1)

            total_page_number_element = str(self.driver.find_element_by_xpath(total_page_number_xpath).text)



            if total_page_number_element != '':

                self.total_page_number = int(total_page_number_element)

            else:

                time.sleep(1)



        self.total_page_number = int(total_page_number_element) 



        print("Total Page: " + str(self.total_page_number))





    def Item_Number_filling(self):

        self.driver.get("http://aci.customs.gov.tw/APIM/IE07?opener=true&?cust_Cd=AW")

        time.sleep(1)

        self.filling_application_number()

        self.checking_status_Message()

        self.pressing_searching_button_F6()

        self.finding_total_page_number()

        item_number_count = 1

        page_number_count = 0

        while page_number_count < self.total_page_number:#開始檢查頁面

            #等待資料傳輸完成

            self.checking_status_Message()  



            for number in range(1 , 21):#開始檢查每頁的每20個貨品明細





                #納稅辦法的Xpath(絕對位置)

                taxPayMethodR_Xpath = "/html/body/div[2]/div[3]/form/div[1]/div[1]/table/tbody/tr[3]/td/div[2]/table/tbody/tr[{}]/td[2]/select".format(number * 2 + 1)



                #找到下拉式選單

                check_element_caught = ''

                

                while check_element_caught == '':

                    #.sleep(0.75)

                    check_element_caught = Select(self.driver.find_element_by_xpath(taxPayMethodR_Xpath))



                #讀取選單內的以選取文字內容

                check_element_existis = ''

                

                while check_element_existis == '':

                    #time.sleep(0.75)

                    check_element_existis = str(check_element_caught.first_selected_option.text)

                    if check_element_existis == '':

                        time.sleep(0.5)

                    if check_element_existis == '':

                        break



                print(check_element_existis)

                if check_element_existis[0:2] == "92":

                    self.item_number_list.append(item_number_count)



                item_number_count = item_number_count + 1





            page_number_count = page_number_count + 1



            if page_number_count < self.total_page_number:

                self.pressing_next_item_buttom()





    def store_ware_finding(self):#ID21

        self.driver.get("http://aci.customs.gov.tw/APIM/ID21?opener=true&?cust_Cd=AW")

        time.sleep(1)

        self.filling_application_number()



        #self.driver.execute_script("f6();")

        self.pressing_searching_button_F6()

        self.checking_status_Message()

        store_ware = ""

        store_ware = str(self.driver.execute_script("return $('#storWareCdDesc').val();"))

        while store_ware == "":

            time.sleep(1)

            store_ware = str(self.driver.execute_script("return $('#storWareCdDesc').val();"))

        print(store_ware)



        self.store_ware = store_ware



    def item_detail_finding(self):#ID24

        self.driver.get("http://aci.customs.gov.tw/APIM/ID24?opener=true&?cust_Cd=AW")

        Iterat_index_item_number_list = 0

        length = len(self.item_number_list)

        itemNo_temp = 0

        self.filling_application_number()

        #item_number_list

        for number in self.item_number_list:

            self.filling_item_number(number)

            self.pressing_searching_button_F6()

            self.checking_status_Message()





            #取得重量

            qty_current = self.collecting_element_by_javascript_commend("qty")

            print(qty_current)

            self.item_details_Unit.append(qty_current)



            #取得淨重

            netWeight_current = self.collecting_element_by_javascript_commend("netWeight")

            print(netWeight_current)

            self.item_details_KGM.append(netWeight_current)



            #取得數量單位

            qtyUnit_current = self.collecting_element_by_finding_id("qtyUnit")

            print(qtyUnit_current)

            self.item_details_qtyUnit.append(qtyUnit_current)

            



            





            #document.querySelector("#qtyUnit")

            Iterat_index_item_number_list = Iterat_index_item_number_list + 1



    def generating_time_format(self):

        now_time = str(datetime.datetime.now())



        self.year = int(now_time[0:4])-1911

        self.month = str(now_time[5:7])

        self.day = str(now_time[8:10])



    def filling_qtyUnit(self):

        Iterat_index = 0

        string_qtyUnit = ""

        length = len(self.item_number_list)

        while Iterat_index < length:

            number = self.item_number_list[Iterat_index]

            Unit = self.item_details_Unit[Iterat_index]

            qtyUnit = self.item_details_qtyUnit[Iterat_index][0:3]

            if Iterat_index != length -1:

                string_qtyUnit += "ITEM {}：{} {}\n".format(number, Unit, qtyUnit)

            else:#不換行

                string_qtyUnit += "ITEM {}：{} {}".format(number, Unit, qtyUnit)

            Iterat_index = Iterat_index + 1

        self.doc.tables[0].cell(3, 1).text = string_qtyUnit





    def filling_KGM(self):

        Iterat_index = 0

        str_KGM = ""

        length = len(self.item_number_list)

        while Iterat_index < length:

            number = self.item_number_list[Iterat_index]

            KGM = self.item_details_KGM[Iterat_index]

            if Iterat_index != length -1:

                str_KGM += "ITEM {}：{} KGM\n".format(number, KGM)

            else:#不換行

                str_KGM += "ITEM {}：{} KGM".format(number, KGM)

            Iterat_index = Iterat_index + 1

        self.doc.tables[0].cell(4, 1).text = str_KGM





    def filling_data_in_page(self):

        self.generating_time_format()

        #標題

        Header_Date_Case_Number = "{}年{}月{}日第0{}號".format(self.year, self.month, self.day, self.input_temp_Case_Number)

        self.doc.paragraphs[1].text = Header_Date_Case_Number



        #報單號碼

        Header_Application_Number = "（報單號碼：{})".format(self.application_with_slash())

        self.doc.paragraphs[2].text = Header_Application_Number



        #案件號碼

        Case_number_paragraphs = "逾期貨物或聲明放棄貨物處理記錄\n（{}）年（□滯報□滯納■逾退）（ 五 ）字第0{}號".format(self.year, self.input_temp_Case_Number)

        self.doc.tables[0].cell(0, 1).text = Case_number_paragraphs



        #填寫件數

        self.filling_qtyUnit()



        #填寫重量

        self.filling_KGM()



        #填寫貨品存放位置

        self.doc.tables[0].cell(6, 1).text = "貨物存放____{}____庫____________間_______________區".format(self.store_ware[0:2])





    def save_file(self):

        file_name = "對貨報告{}-0{}.docx".format(self.year, self.input_temp_Case_Number)

        self.doc.save(file_name)



    def open_file(self):



        self.doc = Document('Cargo_check_109_0333.docx')





    def __init__(self):



        self.application_number = ""

        self.store_ware = ""

        self.item_details_KGM = []

        self.item_details_Unit = []

        self.item_details_qtyUnit = []

        self.item_number_list = []

        self.total_page_number = 0



        self.application_number = input("請輸入報單號碼,輸入z結束\n")

        self.input_temp_Case_Number = input("請輸入對貨報告表的案號,共三碼\n")

        



        ID = input("請輸入帳號")

        

        Password = input("請輸入密碼")



        options = webdriver.ChromeOptions()



        prefs = {

            'profile.default_content_setting_values':

                {

                    'notifications': 2

                }

        }



        options.add_experimental_option('prefs', prefs)



        options.add_argument("disable-infobars")  



        # 打啟動selenium，務必確認driver檔案跟python檔案要在同個資料夾中



        driver = webdriver.Chrome(options=options)



        self.driver = driver





        driver.get("http://aci.customs.gov.tw/portal/Login_main")



        time.sleep(0.5)

        self.driver.find_element_by_xpath('/html/body/form/table/tbody/tr[2]/td[3]/div/div[1]/table/tbody/tr[1]/td[2]/font/input[1]').send_keys(ID)




        self.driver.find_element_by_xpath('/html/body/form/table/tbody/tr[2]/td[3]/div/div[1]/table/tbody/tr[2]/td[2]/input').send_keys(Password)



        time.sleep(1)

        self.driver.find_element_by_xpath('/html/body/form/table/tbody/tr[2]/td[3]/div/div[1]/table/tbody/tr[4]/td/input[1]').click()

        

        time.sleep(1)



        self.Item_Number_filling()









if __name__=='__main__':



    First = goods_checking_document_generator()

    First.store_ware_finding()

    First.item_detail_finding()



    First.open_file()

    First.filling_data_in_page()

    First.save_file()



    #打開檔案

    file_name = "對貨報告{}-0{}.docx".format(First.year, First.input_temp_Case_Number)

    file_path = "F:\D\Checking_Good\Overdue_Returning\109\{}".format(file_name)

    os.startfile(file_name)

    print("已生成對貨報告")
