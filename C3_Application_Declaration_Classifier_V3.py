from selenium import webdriver

from selenium.webdriver.support.ui import WebDriverWait

from selenium.webdriver.support.ui import Select

import time

import datetime

import os

import playsound

from docx import Document

from docx.shared import Pt

class IE07():
	def application_number_check(self):
		if len(self.application_number) > 14 or len(self.application_number) < 12:
			return False
		else:
			return True


	def application_number_convert(self):
		if len(self.application_number) == 14:
			return self.application_number
		elif len(self.application_number) == 12:
			return self.application_number[0:2] + "  " + self.application_number[2:14]
		else:
			return ''


	def application_number_transformation(self):
		if self.application_number_check() == True:
			return self.application_number_convert()
		else:
			return ''


	def filling_application_number(self):

		application_number_transformation_finished = self.application_number_transformation() 

		search_input = self.driver.find_element_by_name("declNo1")
		search_input.clear()
		search_input.send_keys(application_number_transformation_finished[0:2])

		search_input = self.driver.find_element_by_name("declNo2")
		search_input.clear()
		search_input.send_keys(application_number_transformation_finished[2:4])

		search_input = self.driver.find_element_by_name("declNo3")
		search_input.clear()
		search_input.send_keys(application_number_transformation_finished[4:6])

		search_input = self.driver.find_element_by_name("declNo4")
		search_input.clear()
		search_input.send_keys(application_number_transformation_finished[6:9])

		search_input = self.driver.find_element_by_name("declNo5")
		search_input.clear()
		search_input.send_keys(application_number_transformation_finished[9:14])


	def pressing_searching_button_F6(self):
		commit = self.driver.find_element_by_name("btn_F6")

		while commit == '':
			time.sleep(0.5)
			commit = self.driver.find_element_by_name("btn_F6")

		commit.click()


	def pressing_next_item_buttom(self):

		click_lock = True
		while click_lock:
			try:
				searching_page_down_button = self.driver.find_element_by_id("rcfw")
				searching_page_down_button.click()
			except Exception as error:
				continue
			click_lock = False


	def finding_total_page_number(self):

		total_page_number_xpath = "/html/body/div[2]/div[3]/form/div[1]/div[2]/table/tbody/tr/td/span[1]/span"
		total_page_number_element = str(self.driver.find_element_by_xpath(total_page_number_xpath).text)

		while total_page_number_element == '':
			time.sleep(1)
			total_page_number_element = str(self.driver.find_element_by_xpath(total_page_number_xpath).text)

			if total_page_number_element != '':
				self.total_page_number = int(total_page_number_element)
			else:
				time.sleep(1)

		self.total_page_number = int(total_page_number_element)	

		print("Total Page: " + str(self.total_page_number))


	def checking_status_Message(self):
		#checking_status_Message

		executable_javascript_commend = "return $(\"input[name='custValAmtR']\").eq({}).val();"	
		time.sleep(1)
		check_empty = ''
		while check_empty == '':
			check_empty = str(self.driver.execute_script(executable_javascript_commend))
			time.sleep(1)


	def collecting_item_details(self):

		self.filling_application_number()

		self.pressing_searching_button_F6()

		self.checking_status_Message()

		self.finding_total_page_number()

		item_number_count = 1
		page_number = 0 
		
		while page_number < self.total_page_number:#開始檢查頁面

			#等待資料傳輸完成
			self.checking_status_Message()

			#是否分到進口一課業務股
			if self.Destination_Export_Cargo_Clearance_Subsection == True or self.duty_processing_cost == True:
				pass


			for number in range(1 , 21):#開始檢查每頁的每20個貨品明細

				if self.Destination_Export_Cargo_Clearance_Subsection == True or self.Destination_Export_Cargo_Clearance_Subsection:
					pass

				#讀取完稅價格
				Dutiable_value_commend = "return $(\"input[name='custValAmtR']\").eq({}).val();".format(number)
				check_element_existis = str(self.driver.execute_script(Dutiable_value_commend)).replace(",", "")

				if check_element_existis == '':
					time.sleep(0.5)
					check_element_existis = str(self.driver.execute_script(Dutiable_value_commend)).replace(",", "")
					
				if check_element_existis == '':
					break
				

				Dutiable_value_maximum_current = check_element_existis
				
				print(Dutiable_value_maximum_current)
				if self.Dutiable_value_maximum < int(Dutiable_value_maximum_current):
					#print("Item: " + str(item_number_count))
					self.Dutiable_value_maximum = int(Dutiable_value_maximum_current)
					#print(self.Dutiable_value_maximum)

					#讀取稅則號別
					executable_javascript_commend = "return $(\"input[name='cccCodeR']\").eq({}).val();".format(number)
					self.cccCode_maximum = str(self.driver.execute_script(executable_javascript_commend))
					#print(self.cccCode_maximum)

					#讀取納稅辦法
					taxPayMethodR_Xpath = "/html/body/div[2]/div[3]/form/div[1]/div[1]/table/tbody/tr[3]/td/div[2]/table/tbody/tr[{}]/td[2]/select".format(number * 2 + 1)

					check_element_caught = ''
					while check_element_caught == '':
						time.sleep(1)
						check_element_caught = Select(self.driver.find_element_by_xpath(taxPayMethodR_Xpath))

					check_element_existis = ''
					while check_element_existis == '':
						time.sleep(1)
						check_element_existis = str(check_element_caught.first_selected_option.text)

					self.Duty_Treatment_maximum = check_element_existis
					#if self.Duty_Treatment_maximum[0:2] == "55" or self.Duty_Treatment_maximum[0:2] == "58" or self.Duty_Treatment_maximum[0:2] == "99" or self.Duty_Treatment_maximum[0:2] == "90":
						#self.Destination_Export_Cargo_Clearance_Subsection = True

					for treatment in self.Special_Duty_Treatment:
						if treatment == self.Duty_Treatment_maximum[0:2]:
							self.Destination_Export_Cargo_Clearance_Subsection = True

					if self.Duty_Treatment_maximum[0:2] == "3F":
						self.duty_processing_cost = True

					#print(self.Duty_Treatment_maximum)

					self.items_max = item_number_count

				item_number_count = item_number_count + 1
				
			page_number = page_number + 1

			if page_number < self.total_page_number:
				self.pressing_next_item_buttom()


	def reset_all_data(self):

		self.items_max = 0

		self.application_number = ""

		self.Dutiable_value_maximum = 0

		self.Duty_Treatment_maximum = ""

		self.cccCode_maximum = ""

		self.total_page_number = 0

		self.Destination_Export_Cargo_Clearance_Subsection = False

		self.duty_processing_cost = False


	def Delivery_destination(self):
		print("Max price: " + str(self.Dutiable_value_maximum))
		print("Max items: "+ str(self.items_max))
		print("Max cccCode: "+ str(self.cccCode_maximum))
		print("Max Duty_Treatment: "+ self.Duty_Treatment_maximum[0:2])
		if self.Duty_Treatment_maximum != "" or self.cccCode_maximum != "":
			if self.Duty_Treatment_maximum[0:2].isnumeric():
				if self.Destination_Export_Cargo_Clearance_Subsection == True:
					Export_Cargo_Clearance_Subsection_destination = int(self.Duty_Treatment_maximum[0:2])
					
					if self.Destination_Export_Cargo_Clearance_Subsection == True:
						for i in range(0, 5):
							playsound.playsound('preview .mp3', True)
						if Export_Cargo_Clearance_Subsection_destination == 55 or Export_Cargo_Clearance_Subsection_destination == 99:
							print("Re-import Goods")
						elif Export_Cargo_Clearance_Subsection_destination == 58:
							print("Duty Treatment 58")
						elif Export_Cargo_Clearance_Subsection_destination == 90:
							print("Trilateral Trade Goods")
						elif Export_Cargo_Clearance_Subsection_destination == 68:
							print("The Temporary Admission")
				else:
					if self.duty_processing_cost != True:
						if self.cccCode_maximum[0:1] == "0":
							Subsection_destination = int(self.cccCode_maximum[1:2])
						else:
							Subsection_destination = int(self.cccCode_maximum[0:2])
						if (Subsection_destination > 0 and Subsection_destination < 22) or Subsection_destination == 98:
							print("First Subsection")
							playsound.playsound('preview .mp3', True)			
						elif Subsection_destination > 21 and Subsection_destination < 35:
							print("Second Subsection")
							for i in range(0, 2):
								playsound.playsound('preview .mp3', True)	
						elif Subsection_destination > 34 and Subsection_destination < 61:
							print("Third Subsection")
							for i in range(0, 3):
								playsound.playsound('preview .mp3', True)				
						elif Subsection_destination > 59 and Subsection_destination < 70:
							print("Fourth Subsection")
							for i in range(0, 4):
								playsound.playsound('preview .mp3', True)
						else:
							print("Wrong Destination")
							for i in range(0, 6):
								playsound.playsound('preview .mp3', True)
			else:
				if self.duty_processing_cost == True:
					print("Fourth Subsection, Duty levied in accordance with processing Cost")
					for i in range(0, 4):
						playsound.playsound('preview .mp3', True)


	def append_application_number(self, List):
		#Application_Number_List.append(self.application_number)
		List.append(self.application_number)


	def __init__(self):

		ID = "010824"
		
		Password = "xup6xu;4Wu/6"

		options = webdriver.ChromeOptions()

		prefs = {
    		'profile.default_content_setting_values':
        		{
            		'notifications': 2
        		}
		}

		options.add_experimental_option('prefs', prefs)

		options.add_argument("disable-infobars")  

		# 打啟動selenium，務必確認driver檔案跟python檔案要在同個資料夾中

		driver = webdriver.Chrome(options=options)


		driver.get("http://aci.customs.gov.tw/portal/Login_main")

		time.sleep(0.5)

		context = driver.find_element_by_css_selector('#userId')

		
		context.send_keys(ID)

		#輸入password

		context = driver.find_element_by_css_selector('#userPwd')
		
		context.send_keys(Password)

		commit = driver.find_element_by_css_selector('#loginByUserPwd')

		commit.click()

		driver.get("http://aci.customs.gov.tw/APIM/IE07?opener=true&?cust_Cd=AW")

		time.sleep(1)

		self.driver = driver

		self.application_number = ""

		self.Dutiable_value_maximum = 0

		self.Duty_Treatment_maximum = -1

		self.cccCode_maximum = ""

		self.total_page_number = 0

		self.Destination_Export_Cargo_Clearance_Subsection = False

		self.Special_Duty_Treatment =[55, 68, 90, 99]

		self.items_max = 0

		self.duty_processing_cost = False


class document_generating():

	def print_Document(self):
		self.doc.save('A.doc')

	def Document_manipulating(self):
		now_time = str(datetime.datetime.now())

		paragraph = self.doc.add_paragraph()

		number_count = 0
		String_to_print = ""
		String_to_print += "報單清表列印時間: "
		String_to_print += now_time[0:19]
		String_to_print += "\n"
		for number in self.List:
			if number_count % 3 == 2:
				String_to_print += number
				String_to_print += "\n"
			else:
				String_to_print += number 
				String_to_print += "\t"
			number_count = number_count +1

		run = paragraph.add_run(String_to_print)
		font = run.font
		font.name = 'Calibri'
		font.size = Pt(20)

	def __init__(self, Application_Number_List):
		document = Document()
		self.doc = document
		self.List = Application_Number_List
		



if __name__=='__main__':

	First = IE07()
	finish_state = True

	Application_Number_List = []


	while finish_state:

		Error_Collecting = False 
		Error_Distributing = False

		First.application_number = input("請輸入報單號碼,輸入z結束\n")
		if First.application_number == "z":
			finish_state = False
			break
		try:
			First.driver.execute_script("sf1Clear();")
			First.collecting_item_details()
		except:
			print("An exception occurred in Collecting !")
			Error_Collecting = True

		#try:
		if Error_Collecting != True:
			First.Delivery_destination()
		#except:
			#print("An exception occurred in Distributing !")
			#Error_Distributing = True

		if Error_Collecting == True or Error_Distributing == True:
			for i in range(0, 6):
				playsound.playsound('preview .mp3', True)
			Error_Collecting = False
			Error_Distributing = False

		First.append_application_number(Application_Number_List)

		First.reset_all_data()

	Second = document_generating(Application_Number_List)
	Second.Document_manipulating()
	Second.print_Document()
